#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word Document Text Extraction Tool for Pyunto Intelligence

This module provides functions to extract text from Microsoft Word documents for
analysis with Pyunto Intelligence.
"""

import os
import sys
import argparse
import logging
import glob
from datetime import datetime
from typing import Dict, List, Optional, Union, Any, Tuple

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

try:
    import docx
except ImportError:
    logger.error("python-docx is required. Install it with 'pip install python-docx'")
    sys.exit(1)

class WordTextExtractor:
    """Word document text extraction tool for Pyunto Intelligence."""
    
    def __init__(self):
        """Initialize the Word document extractor."""
        pass
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from a Word document.
        
        Args:
            docx_path: Path to the Word document file
            
        Returns:
            Extracted text as a string
            
        Raises:
            FileNotFoundError: If the document file doesn't exist
            Exception: If text extraction fails
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Word document not found: {docx_path}")
        
        try:
            logger.info(f"Extracting text from {docx_path}")
            
            # Load the document
            doc = docx.Document(docx_path)
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        # Get text from cell paragraphs
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                row_text.append(paragraph.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            # Combine all text with newlines
            result = "\n\n".join(full_text)
            
            logger.info(f"Successfully extracted {len(result)} characters from {docx_path}")
            return result
            
        except Exception as e:
            logger.error(f"Error extracting text from Word document: {e}")
            
            # Try alternative extraction method
            try:
                logger.info("Trying alternative extraction method...")
                import docx2txt
                alt_text = docx2txt.process(docx_path)
                logger.info(f"Alternative method extracted {len(alt_text)} characters")
                return alt_text
            except ImportError:
                logger.warning("docx2txt not available for fallback extraction")
            except Exception as alt_e:
                logger.error(f"Alternative extraction method also failed: {alt_e}")
            
            raise Exception(f"Could not extract text from {docx_path}: {str(e)}")
    
    def extract_text_and_metadata(self, docx_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from a Word document.
        
        Args:
            docx_path: Path to the Word document file
            
        Returns:
            Tuple containing:
            - Extracted text as a string
            - Dictionary of document metadata
            
        Raises:
            FileNotFoundError: If the document file doesn't exist
            Exception: If extraction fails
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Word document not found: {docx_path}")
        
        try:
            logger.info(f"Extracting text and metadata from {docx_path}")
            
            # Load the document
            doc = docx.Document(docx_path)
            
            # Extract text
            text = self.extract_text_from_docx(docx_path)
            
            # Extract core properties
            core_props = doc.core_properties
            
            # Build metadata dictionary
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "created": core_props.created,
                "modified": core_props.modified,
                "last_modified_by": core_props.last_modified_by,
                "revision": core_props.revision,
                "category": core_props.category,
                "comments": core_props.comments,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "document_type": "docx",
                "word_count": len(text.split()),
                "character_count": len(text),
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting text and metadata: {e}")
            raise
    
    def extract_structured_content(self, docx_path: str) -> Dict[str, List[Dict[str, Any]]]:
        """
        Extract structured content from a Word document.
        
        Args:
            docx_path: Path to the Word document file
            
        Returns:
            Dictionary containing structured content elements:
            - 'headings': List of heading dictionaries with 'level' and 'text'
            - 'paragraphs': List of paragraph texts
            - 'lists': List of list items
            - 'tables': List of table data
            
        Raises:
            FileNotFoundError: If the document file doesn't exist
            Exception: If extraction fails
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Word document not found: {docx_path}")
        
        try:
            logger.info(f"Extracting structured content from {docx_path}")
            
            # Load the document
            doc = docx.Document(docx_path)
            
            # Initialize structured content
            structured_content = {
                "headings": [],
                "paragraphs": [],
                "lists": [],
                "tables": []
            }
            
            # Process paragraphs
            current_list_items = []
            in_list = False
            
            for para in doc.paragraphs:
                if not para.text.strip():
                    continue
                    
                # Check if it's a heading
                if para.style.name.startswith('Heading'):
                    try:
                        level = int(para.style.name.replace('Heading ', ''))
                    except ValueError:
                        level = 1
                        
                    structured_content["headings"].append({
                        "level": level,
                        "text": para.text
                    })
                
                # Check if it's a list item
                elif para.style.name.startswith('List'):
                    current_list_items.append(para.text)
                    in_list = True
                
                # Regular paragraph
                else:
                    # If we were in a list and now we're not, add the list
                    if in_list and current_list_items:
                        structured_content["lists"].append(current_list_items)
                        current_list_items = []
                        in_list = False
                    
                    structured_content["paragraphs"].append(para.text)
            
            # Add any remaining list items
            if current_list_items:
                structured_content["lists"].append(current_list_items)
            
            # Process tables
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Concatenate all paragraphs in the cell
                        cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                        row_data.append(cell_text)
                    table_data.append(row_data)
                
                if table_data:
                    structured_content["tables"].append(table_data)
            
            return structured_content
            
        except Exception as e:
            logger.error(f"Error extracting structured content: {e}")
            raise
    
    def extract_tables(self, docx_path: str) -> List[List[List[str]]]:
        """
        Extract tables from a Word document.
        
        Args:
            docx_path: Path to the Word document file
            
        Returns:
            List of tables, where each table is a list of rows, 
            and each row is a list of cell values
            
        Raises:
            FileNotFoundError: If the document file doesn't exist
            Exception: If extraction fails
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"Word document not found: {docx_path}")
        
        try:
            logger.info(f"Extracting tables from {docx_path}")
            
            # Load the document
            doc = docx.Document(docx_path)
            
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        # Concatenate all paragraphs in the cell
                        cell_text = "\n".join(p.text for p in cell.paragraphs if p.text.strip())
                        row_data.append(cell_text.strip())
                    table_data.append(row_data)
                
                if table_data:
                    tables.append(table_data)
            
            logger.info(f"Extracted {len(tables)} tables from {docx_path}")
            return tables
            
        except Exception as e:
            logger.error(f"Error extracting tables: {e}")
            raise
    
    def save_text(self, text: str, output_path: str) -> None:
        """
        Save extracted text to a file.
        
        Args:
            text: Extracted text
            output_path: Path to save the text file
        """
        # Ensure the output directory exists
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        
        # Write the text to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        logger.info(f"Text saved to {output_path}")
    
    def batch_process(self, input_dir: str, output_dir: str, recursive: bool = False) -> Dict[str, str]:
        """
        Process multiple Word document files in a directory.
        
        Args:
            input_dir: Directory containing Word document files
            output_dir: Directory to save extracted text files
            recursive: Whether to process documents in subdirectories
            
        Returns:
            Dictionary mapping input files to output files
        """
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)
        
        # Find Word document files (both .docx and .doc)
        if recursive:
            docx_pattern = os.path.join(input_dir, "**", "*.docx")
            docx_files = glob.glob(docx_pattern, recursive=True)
            
            doc_pattern = os.path.join(input_dir, "**", "*.doc")
            doc_files = glob.glob(doc_pattern, recursive=True)
        else:
            docx_pattern = os.path.join(input_dir, "*.docx")
            docx_files = glob.glob(docx_pattern)
            
            doc_pattern = os.path.join(input_dir, "*.doc")
            doc_files = glob.glob(doc_pattern)
        
        # Combine file lists
        all_files = docx_files + doc_files
        
        if not all_files:
            logger.warning(f"No Word document files found in {input_dir}")
            return {}
        
        logger.info(f"Found {len(all_files)} Word document files to process")
        
        results = {}
        success_count = 0
        
        # Process each file
        for doc_file in all_files:
            try:
                # Determine output path
                rel_path = os.path.relpath(doc_file, input_dir)
                rel_dir = os.path.dirname(rel_path)
                base_name = os.path.splitext(os.path.basename(doc_file))[0]
                
                # Create subdirectory in output dir if the file is in a subdirectory
                out_subdir = os.path.join(output_dir, rel_dir) if rel_dir != '.' else output_dir
                os.makedirs(out_subdir, exist_ok=True)
                
                output_file = os.path.join(out_subdir, f"{base_name}.txt")
                
                # Extract text and save it
                text = self.extract_text_from_docx(doc_file)
                self.save_text(text, output_file)
                
                results[doc_file] = output_file
                success_count += 1
                
            except Exception as e:
                logger.error(f"Error processing {doc_file}: {e}")
                results[doc_file] = None
        
        logger.info(f"Successfully processed {success_count} of {len(all_files)} files")
        return results


def main():
    """Command-line interface for Word document text extraction."""
    parser = argparse.ArgumentParser(
        description="Extract text from Word documents for Pyunto Intelligence"
    )
    
    # Main arguments
    parser.add_argument(
        "input",
        help="Input Word document file or directory"
    )
    
    parser.add_argument(
        "-o", "--output",
        help="Output text file or directory"
    )
    
    # Extraction options
    parser.add_argument(
        "--metadata",
        action="store_true",
        help="Extract and include document metadata"
    )
    
    parser.add_argument(
        "--structured",
        action="store_true",
        help="Extract structured content (headings, paragraphs, lists, tables)"
    )
    
    parser.add_argument(
        "--tables-only",
        action="store_true",
        help="Extract only tables from the document"
    )
    
    # Batch processing
    parser.add_argument(
        "--batch",
        action="store_true",
        help="Process multiple Word document files in a directory"
    )
    
    parser.add_argument(
        "--output-dir",
        help="Output directory for batch processing"
    )
    
    parser.add_argument(
        "--recursive",
        action="store_true",
        help="Process documents in subdirectories recursively"
    )
    
    # Other options
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Enable verbose logging"
    )
    
    args = parser.parse_args()
    
    # Set logging level based on verbosity
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    try:
        # Initialize extractor
        extractor = WordTextExtractor()
        
        if args.batch or os.path.isdir(args.input):
            # Batch processing mode
            input_dir = args.input
            output_dir = args.output_dir or args.output or "extracted_texts"
            
            results = extractor.batch_process(
                input_dir=input_dir,
                output_dir=output_dir,
                recursive=args.recursive
            )
            
            # Print summary
            success_count = sum(1 for v in results.values() if v is not None)
            logger.info(f"Batch processing complete. Success: {success_count}/{len(results)}")
            
        else:
            # Single file mode
            docx_path = args.input
            
            if not args.output:
                # Generate output path from input
                base_name = os.path.splitext(os.path.basename(docx_path))[0]
                output_path = f"{base_name}.txt"
            else:
                output_path = args.output
            
            # Handle different extraction modes
            if args.metadata:
                # Extract text and metadata
                text, metadata = extractor.extract_text_and_metadata(docx_path)
                
                # Save text
                extractor.save_text(text, output_path)
                
                # Save metadata to a JSON file
                import json
                meta_path = f"{os.path.splitext(output_path)[0]}_metadata.json"
                
                # Convert datetime objects to strings for JSON serialization
                for key, value in metadata.items():
                    if isinstance(value, datetime):
                        metadata[key] = value.isoformat()
                
                with open(meta_path, 'w', encoding='utf-8') as f:
                    json.dump(metadata, f, indent=2)
                
                logger.info(f"Text extraction complete: {docx_path} -> {output_path}")
                logger.info(f"Metadata saved to: {meta_path}")
                
            elif args.structured:
                # Extract structured content
                content = extractor.extract_structured_content(docx_path)
                
                # Save to JSON file
                import json
                with open(output_path, 'w', encoding='utf-8') as f:
                    json.dump(content, f, indent=2)
                
                logger.info(f"Structured content extraction complete: {docx_path} -> {output_path}")
                
            elif args.tables_only:
                # Extract tables
                tables = extractor.extract_tables(docx_path)
                
                # Save tables as text
                with open(output_path, 'w', encoding='utf-8') as f:
                    for i, table in enumerate(tables):
                        f.write(f"Table {i+1}:\n")
                        for row in table:
                            f.write("| " + " | ".join(row) + " |\n")
                        f.write("\n\n")
                
                logger.info(f"Table extraction complete: {docx_path} -> {output_path}")
                
            else:
                # Standard text extraction
                text = extractor.extract_text_from_docx(docx_path)
                extractor.save_text(text, output_path)
                
                logger.info(f"Text extraction complete: {docx_path} -> {output_path}")
    
    except Exception as e:
        logger.error(f"Error: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
